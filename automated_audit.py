"""
Automated Audit Script.

This script performs the following functions:
1. Reads configurations from the `config.py` file.
2. Iterates over a list of repositories and terms to execute a grep command.
3. Processes the grep command output and stores the results in a DataFrame.
4. Generates a Word document for each repository containing the audit results.

Functions:
- insert_into_table: Inserts a grep output line into a DataFrame.
- process_output: Processes the grep command output for a term and repository.
- add_table_borders: Adds borders to a Word table.
- df_to_word: Converts a DataFrame to a Word document.

Ensure the `terms` list contains the terms you want to audit in config.py.
Ensure you have cloned the repositories you want to audit into the directory specified by `repositoryDirectory` in config.py.
"""

import os
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import config

def insert_into_table(line, term, repository, df):
    components = line.split(":", 2)
    if len(components) < 3:
        print(f"Skipping malformed line: {line}")
        return df
    file_path = components[0]
    line_number = components[1]
    line_content = components[2].strip()

    new_row = pd.DataFrame({
        "Repository": [repository], 
        "File Path": [file_path], 
        "Line Number": [line_number], 
        "Line": [line_content], 
        "Term": [term], 
        "Action": [config.defaultAction],
        "Justification": [config.defaultJustification]
    })
    
    return pd.concat([df, new_row], ignore_index=True)

def process_output(command_output, term, repository, df):
    for line in command_output.split('\n'):
        if line:
            df = insert_into_table(line, term, repository, df)
    return df

def add_table_borders(table):
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def df_to_word(df, doc_name, document_heading):
    doc = Document()
    doc.add_paragraph()
    doc.add_heading(document_heading, level=1)

    term_heading = doc.add_heading('terms in Repository:', level=2)
    if df.empty:
        doc.add_paragraph('None', style='List Bullet')
        print("No terms found in the repository")
    else:
        unique_terms = df["Term"].unique()
        for term in unique_terms:
            doc.add_paragraph(term, style='List Bullet')

        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = column_name

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)

        add_table_borders(table)

    doc.save(doc_name)

# Create the directory for the documents if it doesn't exist
document_folder_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), config.documentFolder)
if not os.path.exists(document_folder_path):
    os.makedirs(document_folder_path)

script_dir = os.path.dirname(os.path.abspath(__file__))

# Iterate over repositories in the specified directory
for repository in os.listdir(config.repositoryDirectory):
    if repository == ".gitignore":
        continue  # Skip the .gitignore file
    path_to_repository = os.path.join(config.repositoryDirectory, repository)
    document_heading = config.documentHeadingTemplate.format(repository)
    doc_name = os.path.join(document_folder_path, f"{repository}_audit.docx")

    if not os.path.exists(path_to_repository):
        print("No directory found: " + path_to_repository)
        continue

    df = pd.DataFrame(columns=["Repository", "File Path", "Line Number", "Line", "Term", "Action", "Justification"])

    for term in config.terms:
        print(f"Grepping for term: {term} in repository: {repository}")
        temp_df = pd.DataFrame(columns=["Repository", "File Path", "Line Number", "Line", "Term", "Action", "Justification"])
        command = config.grepCommandTemplate.format(path=path_to_repository, term=term)
        grep_output = os.popen(command).read()
        temp_df = process_output(grep_output, term, repository, temp_df)
        df = pd.concat([df, temp_df], ignore_index=True)
        print(temp_df)

    print("Saving DataFrame to Word document")
    print(df)
    df_to_word(df, doc_name, document_heading)
